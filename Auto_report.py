#coding:utf-8
import MySQLdb
import time
import datetime
import re
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from pylab import matplotlib,mpl
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontManager
from docx.shared import RGBColor
import ch
import sys
from PyQt4 import QtGui,QtCore
import shutil
import os
from PyQt4.QtGui import *

class Report(QtGui.QWidget):
    def __init__(self, parent = None):
        QtGui.QWidget.__init__(self, parent)
        self.setGeometry(800, 300, 800, 600)
        self.setWindowTitle(u'模型工具自动化生成正负面报告')
        self.setWindowIcon(QtGui.QIcon('timg.jpg'))

        background = QtGui.QPixmap('background.jpg')
        palette = QtGui.QPalette()
        palette.setBrush(self.backgroundRole(), QtGui.QBrush(background))  # 添加背景图片
        self.setPalette(palette)

        # self.output = QtGui.QPushButton(u'选择日期', self)
        # self.output.setGeometry(300, 90, 200, 70)

        self.cal = QtGui.QCalendarWidget(self)
        # self.cal.setGridVisible(True)
        self.cal.setGeometry(250, 50, 300, 180)
        self.connect(self.cal, QtCore.SIGNAL('selectionChanged()'), self.select)

        self.output = QtGui.QPushButton(u'一键导出', self)
        self.output.setGeometry(300, 230, 200, 70)
        self.output.clicked.connect(self.button)

        self.back = QtGui.QPushButton(u'后退', self)
        self.back.setGeometry(20, 20, 60, 40)
        self.back.clicked.connect(self.backbutton)

        self.mysql_conn = MySQLdb.connect(host="192.168.30.2", user='root', passwd='zsce188xa5', db='yqht', charset='utf8')
        self.mysql_cursor = self.mysql_conn.cursor()

        self.document = Document()

        self.today = ''

    def select(self):
        self.today = self.cal.selectedDate()
        self.today = str(self.today.toPyDate()).replace("-","")
        print self.today

    def button(self):
        if self.today:
            file_name = QtGui.QFileDialog.getExistingDirectory()
            print file_name
            if file_name:
                self.main(file_name)
        else:
            msg_box = QMessageBox(QMessageBox.Warning, u"提示", u"请在上方日历中选择最近计算日期")
            msg_box.show()
            qe = QEventLoop()  # 阻塞
            qe.exec_()

    def backbutton(self):
        self.close()

    def main(self,file_name):
        ch.set_ch()

        nowtime = time.strftime("%Y%m%d", time.localtime())
        # now_time = datetime.datetime.now()
        # yes_time = now_time + datetime.timedelta(days=-1)
        # yes_time = yes_time.strftime('%Y%m%d')
        # today = '20170329'
        if not self.today:
            self.today = nowtime
        today = self.today
        print today
        # print yes_time


        self.document.add_heading(today + u'正负面计算验证', 0)
        # 医疗文字书写doc
        sql = '''select * from b_pn_work_jjx WHERE CURPOSNAME = "%s" AND CALINDUSTRY = "3"'''%today
        self.mysql_cursor.execute(sql)
        result = self.mysql_cursor.fetchall()[0]
        # print  result
        number = result[1]
        number = int(number) - 2
        # print number, type(number)
        number = str(number)
        self.document.add_heading(number + u'医疗正负面验证', 1)

        divide = re.search(u'(^.*?中性：\d+)[\s\S](.*?)[\s\S](人工负面-机器正面：\d+ 比例：\d+.\d+)[\s\S](人工[\s\S]*?)[\s\S](人工[\s\S]*?)[\s\S](改进秘书判断：)[\s\S](人工负面-机器正面：\d+ 比例：\d+.\d+)[\s\S](人工[\s\S]*?)[\s\S](人工[\s\S]*?)[\s\S](改进判断整体准确率[\S\s]*?)[\s\S](准确率=[\s\S]*?\+准确率\))',result[-1])
        if divide:
            print divide.group(1),'********',divide.group(2),'******',divide.group(3),'******',divide.group(4),'******',divide.group(5),'******',divide.group(6),'******',divide.group(7),'******',divide.group(8),'******',divide.group(9),'******',divide.group(10),'******',divide.group(11)
            content1 = divide.group(1)
            content2 = divide.group(2)
            content3 = divide.group(3)
            content4 = divide.group(4)
            content5 = divide.group(5)
            content6 = divide.group(6)
            content7 = divide.group(7)
            content8 = divide.group(8)
            content9 = divide.group(9)
            content10 = divide.group(10)
            content11 = divide.group(11)

        write1 = self.document.add_paragraph()
        run = write1.add_run(content1).bold = True

        write6 = self.document.add_paragraph()
        write6 = write6.add_run(content6)
        write6.font.name = 'Estrangelo Edessa'
        write6.font.size = Pt(10)
        write6.font.color.rgb = RGBColor(0,0,255)

        write7 = self.document.add_paragraph()
        write7 = write7.add_run(content7)
        write7.font.name = 'Estrangelo Edessa'
        write7.font.size = Pt(12)

        write8 = self.document.add_paragraph()
        write8 = write8.add_run(content8)
        write8.font.name = 'Estrangelo Edessa'
        write8.font.size = Pt(12)
        write8.font.color.rgb = RGBColor(220,20,60)

        write9 = self.document.add_paragraph()
        write9 = write9.add_run(content9)
        write9.font.name = 'Estrangelo Edessa'
        write9.font.size = Pt(12)

        write10 = self.document.add_paragraph()
        write10 = write10.add_run(content10)
        write10.font.name = 'Estrangelo Edessa'
        write10.font.size = Pt(12)
        write10.font.color.rgb = RGBColor(220,20,60)

        write11 = self.document.add_paragraph()
        write11 = write11.add_run(content11)
        write11.font.name = 'Estrangelo Edessa'
        write11.font.size = Pt(12)



        # 政府文字书写doc
        sql = '''select * from b_pn_work_jjx WHERE CURPOSNAME = "%s" AND CALINDUSTRY = "1"'''%today
        self.mysql_cursor.execute(sql)
        result = self.mysql_cursor.fetchall()[0]
        # print  result
        number = result[1]
        number = int(number) - 2
        # print number, type(number)
        number = str(number)
        self.document.add_heading(number + u'政府正负面验证', 1)

        print result[-1]
        divide = re.search(u'(^.*?中性：\d+)[\s\S](.*?)[\s\S](人工负面-机器正面：\d+ 比例：\d+.\d+)[\s\S](人工[\s\S]*?)[\s\S](人工[\s\S]*?)[\s\S](改进秘书判断：)[\s\S](人工负面-机器正面：\d+ 比例：\d+.\d+)[\s\S](人工[\s\S]*?)[\s\S](人工[\s\S]*?)[\s\S](改进判断整体准确率[\S\s]*?)[\s\S](准确率=[\s\S]*?\+准确率\))',result[-1])
        if divide:
            print divide.group(1),'********',divide.group(2),'******',divide.group(3),'******',divide.group(4),'******',divide.group(5),'******',divide.group(6),'******',divide.group(7),'******',divide.group(8),'******',divide.group(9),'******',divide.group(10),'******',divide.group(11)
            content1 = divide.group(1)
            content2 = divide.group(2)
            content3 = divide.group(3)
            content4 = divide.group(4)
            content5 = divide.group(5)
            content6 = divide.group(6)
            content7 = divide.group(7)
            content8 = divide.group(8)
            content9 = divide.group(9)
            content10 = divide.group(10)
            content11 = divide.group(11)

        write1 = self.document.add_paragraph()
        run = write1.add_run(content1).bold = True

        write6 = self.document.add_paragraph()
        write6 = write6.add_run(content6)
        write6.font.name = 'Estrangelo Edessa'
        write6.font.size = Pt(12)
        write6.font.color.rgb = RGBColor(0,0,255)

        write7 = self.document.add_paragraph()
        write7 = write7.add_run(content7)
        write7.font.name = 'Estrangelo Edessa'
        write7.font.size = Pt(12)

        write8 = self.document.add_paragraph()
        write8 = write8.add_run(content8)
        write8.font.name = 'Estrangelo Edessa'
        write8.font.size = Pt(12)
        write8.font.color.rgb = RGBColor(220,20,60)

        write9 = self.document.add_paragraph()
        write9 = write9.add_run(content9)
        write9.font.name = 'Estrangelo Edessa'
        write9.font.size = Pt(12)

        write10 = self.document.add_paragraph()
        write10 = write10.add_run(content10)
        write10.font.name = 'Estrangelo Edessa'
        write10.font.size = Pt(12)
        write10.font.color.rgb = RGBColor(220,20,60)

        write11 = self.document.add_paragraph()
        write11 = write11.add_run(content11)
        write11.font.name = 'Estrangelo Edessa'
        write11.font.size = Pt(12)

        result = self.get_all()
        # draw_pic1(result)
        self.draw_pic2(result)
        # draw_pic3(result)
        self.draw_pic4(result)
        # draw_pic5(result)
        self.draw_pic6(result)

        name = '%s.docx'%today
        self.document.save(name)
        print sys.path[0],'************',file_name
        shutil.move(sys.path[0]+'\\'+name,file_name+'\\'+name)
        os.remove(sys.path[0] + '\\2.png')
        os.remove(sys.path[0] + '\\4.png')
        os.remove(sys.path[0] + '\\6.png')
        return

    def get_all(self):
        sql_all = '''SELECT * FROM b_pn_work_jjx'''
        self.mysql_cursor.execute(sql_all)
        result = self.mysql_cursor.fetchall()
        print  len(result)
        return result

    def draw_pic2(self,result):
        pic1_61 = []
        pic1_3 = []
        for each in result:
            if each[6] == '1':
                pic1_61.append((str(each[4]),each[59]))
            else:
                pic1_3.append((str(each[4]), each[59]))
        pic1_61.sort(key=lambda x: (x[0], x[0]))
        pic1_61.insert(0,(u'标准','0.8788462'))
        pic1_3.sort(key=lambda x: (x[0], x[0]))
        pic1_3.insert(0,(u'标准','0.4939759'))
        print pic1_61
        print pic1_3
        y1 = [a[1] for a in pic1_61]
        y2 = [b[1] for b in pic1_3]
        x1 = [c for c in range(0,len(y1))]
        x2 = x1
        print x1
        print y1
        print y2
        plt.figure('2.png')
        plt.plot(x1, y1, label=u'政府', linewidth=3, color='r', marker='o',markerfacecolor='blue', markersize=12)
        plt.plot(x2, y2, label=u'医疗')
        plt.xlabel(u'时间')
        plt.ylabel(u'比率')
        plt.title(u'改进秘书判断：负面召回率')
        plt.legend()
        # plt.show()
        plt.savefig('2.png')
        self.document.add_picture('2.png', width=Inches(6.4))


        pic1_6 = pic1_61[-8:]
        pic1_6.insert(0, (u'标准', '0.8788462'))
        pic1_3 = pic1_3[-8:]
        pic1_3.insert(0, (u'标准', '0.4939759'))
        y1 = [a[1] for a in pic1_6]
        y2 = [b[1] for b in pic1_3]
        x1 = [c for c in range(len(pic1_61)-7, len(pic1_61)+1)]
        x1.insert(0,0)
        x2 = x1

        table = self.document.add_table(rows=4, cols=len(y1) + 1)
        for i in range(1,len(y1)):
            table.cell(0, i).text = str(x1[i-1])
        table.cell(0, len(y2)).text = str(x1[-1])
        for i in range(1,len(y1)):
            table.cell(1, i).text = pic1_6[i-1][0].replace("2017","")[0:2]+"."+pic1_6[i-1][0].replace("2017","")[2:]
        table.cell(1, len(y2)).text = pic1_3[len(y2) - 1][0].replace("2017","")[0:2]+"."+pic1_3[len(y2) - 1][0].replace("2017","")[2:]

        for i in range(1,len(y1)):
            table.cell(2, i).text = ['%.3f'%float(pic1_6[i-1][1])]
        print '%.3f'%float(pic1_6[len(y1)-1][1]),len(y1)
        table.cell(2, len(y1)).text = ['%.3f'%float(pic1_6[len(y1)-1][1])]
        for i in range(1,len(y2)):
            table.cell(3, i).text = ['%.3f'%float(pic1_3[i-1][1])]
        table.cell(3, len(y2)).text = ['%.3f'%float(pic1_3[len(y2) - 1][1])]

        table.cell(0, 0).text = u'N'
        table.cell(1, 0).text = u'日期'
        table.cell(2, 0).text = u'政府'
        table.cell(3, 0).text = u'医疗'

    def draw_pic4(self,result):
        pic1_31 = []
        for each in result:
            if each[6] == '3':
                pic1_31.append((str(each[4]),each[67],each[68],each[69]))
            else:
                continue

        pic1_31.sort(key=lambda x: (x[0], x[0]))
        pic1_31.insert(0,(u'标准','0.75431','0.50578','0.605536'))
        # print pic1_31
        y1 = [a[1] for a in pic1_31]
        y2 = [a[2] for a in pic1_31]
        y3 = [a[3] for a in pic1_31]
        x1 = [c for c in range(0,len(y1))]
        print x1
        print y1
        print y2
        print y3
        plt.figure('4.png')
        plt.plot(x1, y1, label=u'整体准确率', linewidth=3, color='r', marker='o',markerfacecolor='blue', markersize=12)
        plt.plot(x1, y2, label=u'整体召回率')
        plt.plot(x1, y3, label=u'F1')
        plt.xlabel(u'时间')
        plt.ylabel(u'比率')
        plt.title(u'改进秘书判断：医疗结果')
        plt.legend()
        # plt.show()
        plt.savefig('4.png')
        self.document.add_picture('4.png', width=Inches(6.35))

        pic1_3 = pic1_31[-8:]
        pic1_3.insert(0, (u'标准', '0.75431', '0.50578', '0.605536'))
        y1 = [a[1] for a in pic1_3]
        y2 = [a[2] for a in pic1_3]
        y3 = [a[3] for a in pic1_3]
        x1 = [c for c in range(len(pic1_31)-7, len(pic1_31)+1)]
        x1.insert(0,0)

        table = self.document.add_table(rows=5, cols=len(y1)+1)
        for i in range(1,len(y1)):
            table.cell(0, i).text = str(x1[i-1])
        table.cell(0, len(y2)).text = str(x1[-1])
        for i in range(1,len(y1)):
            table.cell(1, i).text = pic1_3[i-1][0].replace("2017","")[0:2]+"."+pic1_3[i-1][0].replace("2017","")[2:]
        table.cell(1, len(y2)).text = pic1_3[len(y2) - 1][0].replace("2017","")[0:2]+"."+ pic1_3[len(y2) - 1][0].replace("2017","")[2:]


        for i in range(1,len(y1)):
            table.cell(2, i).text = ['%.3f'%float(pic1_3[i-1][1])]
        table.cell(2, len(y1)).text = ['%.3f'%float(pic1_3[len(y1)-1][1])]
        for i in range(1,len(y2)):
            table.cell(3, i).text = ['%.3f'%float(pic1_3[i-1][2])]
        table.cell(3, len(y2)).text = ['%.3f'%float(pic1_3[len(y2) - 1][2])]
        for i in range(1,len(y2)):
            table.cell(4, i).text = ['%.3f'%float(pic1_3[i-1][3])]
        table.cell(4, len(y2)).text = ['%.3f'%float(pic1_3[len(y2) - 1][3])]

        table.cell(0, 0).text = u'N'
        table.cell(1, 0).text = u'日期'
        table.cell(2, 0).text = u'整体准确率'
        table.cell(3, 0).text = u'整体召回率'
        table.cell(4, 0).text = u'F1'

    def draw_pic6(self,result):
        pic1_31 = []
        for each in result:
            if each[6] == '1':
                pic1_31.append((str(each[4]),each[67],each[68],each[69]))
            else:
                continue

        pic1_31.sort(key=lambda x: (x[0], x[0]))
        pic1_31.insert(0,(u'标准','0.84014','0.8071749','0.8233276'))
        # print pic1_31
        y1 = [a[1] for a in pic1_31]
        y2 = [a[2] for a in pic1_31]
        y3 = [a[3] for a in pic1_31]
        x1 = [c for c in range(0,len(y1))]
        print x1
        print y1
        print y2
        print y3
        plt.figure('6.png')
        plt.plot(x1, y1, label=u'整体准确率', linewidth=3, color='r', marker='o',markerfacecolor='blue', markersize=12)
        plt.plot(x1, y2, label=u'整体召回率')
        plt.plot(x1, y3, label=u'F1')
        plt.xlabel(u'时间')
        plt.ylabel(u'比率')
        plt.title(u'改进秘书判断：政府结果')
        plt.legend()
        # plt.show()
        plt.savefig('6.png')
        self.document.add_picture('6.png', width=Inches(6.35))

        pic1_3 = pic1_31[-8:]
        pic1_3.insert(0, (u'标准', '0.84014', '0.8071749', '0.8233276'))
        y1 = [a[1] for a in pic1_3]
        y2 = [a[2] for a in pic1_3]
        y3 = [a[3] for a in pic1_3]
        x1 = [c for c in range(len(pic1_31)-7, len(pic1_31)+1)]
        x1.insert(0,0)


        table = self.document.add_table(rows=5, cols=len(y1)+1)
        for i in range(1,len(y1)):
            table.cell(0, i).text = str(x1[i-1])
        table.cell(0, len(y2)).text = str(x1[-1])
        for i in range(1,len(y1)):
            table.cell(1, i).text = pic1_3[i-1][0].replace("2017","")[0:2]+"."+pic1_3[i-1][0].replace("2017","")[2:]
        table.cell(1, len(y2)).text = pic1_3[len(y2) - 1][0].replace("2017","")[0:2]+"."+pic1_3[len(y2) - 1][0].replace("2017","")[2:]
        # print pic1_3[len(y2) - 1][0].replace("2017","")[0:2]+"."+pic1_3[len(y2) - 1][0].replace("2017","")[2:]

        for i in range(1,len(y1)):
            table.cell(2, i).text = ['%.3f'%float(pic1_3[i-1][1])]
        table.cell(2, len(y1)).text = ['%.3f'%float(pic1_3[len(y1)-1][1])]
        for i in range(1,len(y2)):
            table.cell(3, i).text = ['%.3f'%float(pic1_3[i-1][2])]
        table.cell(3, len(y2)).text = ['%.3f'%float(pic1_3[len(y2) - 1][2])]
        for i in range(1,len(y2)):
            table.cell(4, i).text = ['%.3f'%float(pic1_3[i-1][3])]
        table.cell(4, len(y2)).text = ['%.3f'%float(pic1_3[len(y2) - 1][3])]

        table.cell(0, 0).text = u'N'
        table.cell(1, 0).text = u'日期'
        table.cell(2, 0).text = u'整体准确率'
        table.cell(3, 0).text = u'整体召回率'
        table.cell(4, 0).text = u'F1'

if __name__ == '__main__':
    app = QtGui.QApplication(sys.argv)
    home = Report()
    home.show()
    sys.exit(app.exec_())